#!/usr/bin/env python3
"""
Backend API test suite for OmniParse after Python 3.13 migration + dependency upgrades.

Tests regression across ALL formats after numpy 2.x, opencv 5.x, onnxruntime 1.27, pillow 12.x upgrade.

Tests:
1. Health and readiness endpoints
2. Auth token minting (mock JWT)
3. Image OCR: PNG and JPG with text
4. Office docs: DOCX and XLSX
5. Scanned PDF OCR: German+English image-only PDF with PaddleOCR
6. Plain text/CSV regression
"""
import io
import json
import os
import sys
import requests
from PIL import Image, ImageDraw, ImageFont

# Backend URL - use localhost as backend runs internally on port 8001
# External URL may not be accessible from within the container
BASE_URL = "http://localhost:8001"

def test_health():
    """Test GET /api/health endpoint."""
    print("\n=== TEST 1: Health Check ===")
    try:
        response = requests.get(f"{BASE_URL}/api/health", timeout=10)
        print(f"Status Code: {response.status_code}")
        print(f"Response: {response.json()}")
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert data.get("status") == "healthy", f"Expected status 'healthy', got {data.get('status')}"
        print("✅ Health check PASSED")
        return True
    except Exception as e:
        print(f"❌ Health check FAILED: {e}")
        return False

def test_ready():
    """Test GET /api/ready endpoint."""
    print("\n=== TEST 2: Readiness Check ===")
    try:
        response = requests.get(f"{BASE_URL}/api/ready", timeout=10)
        print(f"Status Code: {response.status_code}")
        print(f"Response: {response.json()}")
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert data.get("status") == "ready", f"Expected status 'ready', got {data.get('status')}"
        checks = data.get("checks", {})
        assert checks.get("database") == "ok", f"Database check failed: {checks.get('database')}"
        print("✅ Readiness check PASSED")
        return True
    except Exception as e:
        print(f"❌ Readiness check FAILED: {e}")
        return False

def mint_token():
    """Mint a mock JWT token for authentication."""
    print("\n=== TEST 3: Auth Token Minting ===")
    try:
        response = requests.post(
            f"{BASE_URL}/api/v1/auth/token",
            json={"sub": "tester", "roles": ["extractor", "admin"]},
            timeout=10
        )
        print(f"Status Code: {response.status_code}")
        print(f"Response: {response.json()}")
        
        assert response.status_code == 200, f"Expected 200, got {response.status_code}"
        data = response.json()
        assert "accessToken" in data, "No accessToken in response"
        token = data["accessToken"]
        print(f"✅ Token minted successfully: {token[:50]}...")
        return token
    except Exception as e:
        print(f"❌ Token minting FAILED: {e}")
        return None

def create_test_image_with_text(text, filename, format='PNG'):
    """Create a test image with text for OCR testing."""
    # Create a white image
    img = Image.new('RGB', (400, 200), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font, fallback to basic if not available
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except:
        font = ImageFont.load_default()
    
    # Draw text
    draw.text((20, 80), text, fill='black', font=font)
    
    # Save to bytes
    img_bytes = io.BytesIO()
    img.save(img_bytes, format=format)
    img_bytes.seek(0)
    
    return img_bytes

def test_image_ocr_png(token):
    """Test OCR extraction with PNG image containing text."""
    print("\n=== TEST 4: Image OCR - PNG ===")
    if not token:
        print("❌ Skipping - no auth token")
        return False
    
    try:
        # Create a PNG with text
        test_text = "Hello World 2025\nDocument Intelligence Test"
        img_bytes = create_test_image_with_text(test_text, "test.png", "PNG")
        
        files = {
            'file': ('test_ocr.png', img_bytes, 'image/png')
        }
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        print(f"Uploading PNG image with text...")
        response = requests.post(
            f"{BASE_URL}/api/v1/extract/sync",
            files=files,
            headers=headers,
            timeout=60
        )
        
        print(f"Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"Response: {response.text}")
            assert False, f"Expected 200, got {response.status_code}"
        
        data = response.json()
        result = data["result"]
        
        # Critical checks for OCR
        print(f"\n--- OCR Verification ---")
        print(f"ocrUsed: {result.get('ocrUsed')}")
        print(f"ocrEngine: {result.get('ocrEngine')}")
        
        assert result.get("ocrUsed") == True, f"Expected ocrUsed=true, got {result.get('ocrUsed')}"
        
        # Accept both vision API and PaddleOCR (vision is preferred when OCR_VISION=true)
        ocr_engine = result.get("ocrEngine", "")
        assert ocr_engine in ["paddleocr:latin", "vision:openai:gpt-4o", "vision:openai:gpt-4o-mini"], \
            f"Unexpected ocrEngine: '{ocr_engine}'"
        
        markdown = result.get("markdown", "")
        print(f"Markdown length: {len(markdown)} chars")
        print(f"Markdown preview: {markdown[:200]}")
        
        assert len(markdown) > 0, "Markdown is empty"
        print(f"✅ PNG OCR test PASSED (using {ocr_engine})")
        return True
        
    except Exception as e:
        print(f"❌ PNG OCR test FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_image_ocr_jpg(token):
    """Test OCR extraction with JPG image containing text."""
    print("\n=== TEST 5: Image OCR - JPG ===")
    if not token:
        print("❌ Skipping - no auth token")
        return False
    
    try:
        # Create a JPG with text
        test_text = "Invoice 2025\nTotal: $1,234.56"
        img_bytes = create_test_image_with_text(test_text, "test.jpg", "JPEG")
        
        files = {
            'file': ('test_ocr.jpg', img_bytes, 'image/jpeg')
        }
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        print(f"Uploading JPG image with text...")
        response = requests.post(
            f"{BASE_URL}/api/v1/extract/sync",
            files=files,
            headers=headers,
            timeout=60
        )
        
        print(f"Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"Response: {response.text}")
            assert False, f"Expected 200, got {response.status_code}"
        
        data = response.json()
        result = data["result"]
        
        # Critical checks for OCR
        print(f"\n--- OCR Verification ---")
        print(f"ocrUsed: {result.get('ocrUsed')}")
        print(f"ocrEngine: {result.get('ocrEngine')}")
        
        assert result.get("ocrUsed") == True, f"Expected ocrUsed=true, got {result.get('ocrUsed')}"
        
        # Accept both vision API and PaddleOCR (vision is preferred when OCR_VISION=true)
        ocr_engine = result.get("ocrEngine", "")
        assert ocr_engine in ["paddleocr:latin", "vision:openai:gpt-4o", "vision:openai:gpt-4o-mini"], \
            f"Unexpected ocrEngine: '{ocr_engine}'"
        
        markdown = result.get("markdown", "")
        print(f"Markdown length: {len(markdown)} chars")
        print(f"Markdown preview: {markdown[:200]}")
        
        assert len(markdown) > 0, "Markdown is empty"
        print(f"✅ JPG OCR test PASSED (using {ocr_engine})")
        return True
        
    except Exception as e:
        print(f"❌ JPG OCR test FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_test_docx():
    """Create a simple DOCX file for testing."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for omniparse regression testing.')
        doc.add_paragraph('It contains multiple paragraphs.')
        doc.add_heading('Section 1', level=1)
        doc.add_paragraph('Some content in section 1.')
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes
    except ImportError:
        # Fallback: create a minimal valid DOCX (ZIP with required XML)
        import zipfile
        docx_bytes = io.BytesIO()
        with zipfile.ZipFile(docx_bytes, 'w', zipfile.ZIP_DEFLATED) as docx:
            # Minimal DOCX structure
            docx.writestr('[Content_Types].xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>''')
            docx.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>''')
            docx.writestr('word/document.xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>
<w:p><w:r><w:t>Test Document for Omniparse</w:t></w:r></w:p>
<w:p><w:r><w:t>This is a regression test document.</w:t></w:r></w:p>
</w:body>
</w:document>''')
        docx_bytes.seek(0)
        return docx_bytes

def test_office_docx(token):
    """Test extraction with DOCX office document."""
    print("\n=== TEST 6: Office Document - DOCX ===")
    if not token:
        print("❌ Skipping - no auth token")
        return False
    
    try:
        docx_bytes = create_test_docx()
        
        files = {
            'file': ('test_document.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        print(f"Uploading DOCX document...")
        response = requests.post(
            f"{BASE_URL}/api/v1/extract/sync",
            files=files,
            headers=headers,
            timeout=60
        )
        
        print(f"Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"Response: {response.text}")
            assert False, f"Expected 200, got {response.status_code}"
        
        data = response.json()
        result = data["result"]
        
        # Check that OCR was NOT used (office docs don't need OCR)
        print(f"\n--- Document Processing Verification ---")
        print(f"ocrUsed: {result.get('ocrUsed')}")
        
        assert result.get("ocrUsed") == False, f"Expected ocrUsed=false for DOCX, got {result.get('ocrUsed')}"
        
        # Check for structured output
        assert "markdown" in result, "No 'markdown' field in result"
        assert "structured" in result, "No 'structured' field in result"
        
        markdown = result.get("markdown", "")
        structured = result.get("structured", {})
        
        print(f"Markdown length: {len(markdown)} chars")
        print(f"Structured data keys: {list(structured.keys()) if isinstance(structured, dict) else 'not a dict'}")
        
        assert len(markdown) > 0, "Markdown is empty"
        print("✅ DOCX test PASSED")
        return True
        
    except Exception as e:
        print(f"❌ DOCX test FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_test_xlsx():
    """Create a simple XLSX file for testing."""
    try:
        from openpyxl import Workbook
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        
        # Add headers
        ws['A1'] = 'Name'
        ws['B1'] = 'Age'
        ws['C1'] = 'City'
        
        # Add data
        ws['A2'] = 'Alice Johnson'
        ws['B2'] = 28
        ws['C2'] = 'Seattle'
        
        ws['A3'] = 'Bob Smith'
        ws['B3'] = 35
        ws['C3'] = 'Portland'
        
        xlsx_bytes = io.BytesIO()
        wb.save(xlsx_bytes)
        xlsx_bytes.seek(0)
        return xlsx_bytes
    except ImportError:
        # Fallback: create a minimal valid XLSX (ZIP with required XML)
        import zipfile
        xlsx_bytes = io.BytesIO()
        with zipfile.ZipFile(xlsx_bytes, 'w', zipfile.ZIP_DEFLATED) as xlsx:
            # Minimal XLSX structure
            xlsx.writestr('[Content_Types].xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>''')
            xlsx.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>''')
            xlsx.writestr('xl/workbook.xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<sheets>
<sheet name="Sheet1" sheetId="1" r:id="rId1"/>
</sheets>
</workbook>''')
            xlsx.writestr('xl/_rels/workbook.xml.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>''')
            xlsx.writestr('xl/worksheets/sheet1.xml', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
<sheetData>
<row r="1">
<c r="A1" t="inlineStr"><is><t>Name</t></is></c>
<c r="B1" t="inlineStr"><is><t>Age</t></is></c>
<c r="C1" t="inlineStr"><is><t>City</t></is></c>
</row>
<row r="2">
<c r="A2" t="inlineStr"><is><t>Alice</t></is></c>
<c r="B2" t="inlineStr"><is><t>28</t></is></c>
<c r="C2" t="inlineStr"><is><t>Seattle</t></is></c>
</row>
</sheetData>
</worksheet>''')
        xlsx_bytes.seek(0)
        return xlsx_bytes

def test_office_xlsx(token):
    """Test extraction with XLSX office document."""
    print("\n=== TEST 7: Office Document - XLSX ===")
    if not token:
        print("❌ Skipping - no auth token")
        return False
    
    try:
        xlsx_bytes = create_test_xlsx()
        
        files = {
            'file': ('test_spreadsheet.xlsx', xlsx_bytes, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        }
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        print(f"Uploading XLSX document...")
        response = requests.post(
            f"{BASE_URL}/api/v1/extract/sync",
            files=files,
            headers=headers,
            timeout=60
        )
        
        print(f"Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"Response: {response.text}")
            assert False, f"Expected 200, got {response.status_code}"
        
        data = response.json()
        result = data["result"]
        
        # Check that OCR was NOT used (office docs don't need OCR)
        print(f"\n--- Document Processing Verification ---")
        print(f"ocrUsed: {result.get('ocrUsed')}")
        
        assert result.get("ocrUsed") == False, f"Expected ocrUsed=false for XLSX, got {result.get('ocrUsed')}"
        
        # Check for structured output
        assert "markdown" in result, "No 'markdown' field in result"
        assert "structured" in result, "No 'structured' field in result"
        
        markdown = result.get("markdown", "")
        structured = result.get("structured", {})
        
        print(f"Markdown length: {len(markdown)} chars")
        print(f"Structured data keys: {list(structured.keys()) if isinstance(structured, dict) else 'not a dict'}")
        
        assert len(markdown) > 0, "Markdown is empty"
        print("✅ XLSX test PASSED")
        return True
        
    except Exception as e:
        print(f"❌ XLSX test FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_ocr_german_english_pdf(token):
    """Test OCR extraction with German+English image-only PDF (CRITICAL TEST)."""
    print("\n=== TEST 8: OCR Test - Scanned PDF (German+English) [CRITICAL] ===")
    if not token:
        print("❌ Skipping - no auth token")
        return False
    
    try:
        pdf_path = "/app/de_scan.pdf"
        
        with open(pdf_path, 'rb') as f:
            files = {
                'file': ('de_scan.pdf', f, 'application/pdf')
            }
            headers = {
                'Authorization': f'Bearer {token}'
            }
            
            print(f"Uploading {pdf_path}...")
            response = requests.post(
                f"{BASE_URL}/api/v1/extract/sync",
                files=files,
                headers=headers,
                timeout=120  # OCR can take longer
            )
        
        print(f"Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"Response: {response.text}")
            assert False, f"Expected 200, got {response.status_code}"
        
        data = response.json()
        result = data["result"]
        
        # Critical checks for OCR
        print(f"\n--- OCR Verification ---")
        print(f"ocrUsed: {result.get('ocrUsed')}")
        print(f"ocrEngine: {result.get('ocrEngine')}")
        
        assert result.get("ocrUsed") == True, f"Expected ocrUsed=true, got {result.get('ocrUsed')}"
        assert result.get("ocrEngine") == "paddleocr:latin", f"Expected ocrEngine='paddleocr:latin', got '{result.get('ocrEngine')}'"
        
        # Check markdown content
        markdown = result.get("markdown", "")
        print(f"\nMarkdown length: {len(markdown)} chars")
        print(f"Markdown preview (first 500 chars):\n{markdown[:500]}")
        
        # Verify German text with umlauts is present
        german_words = ["München", "Straße", "schön", "Müller"]
        found_german = []
        for word in german_words:
            if word in markdown:
                found_german.append(word)
        
        print(f"\nGerman words found: {found_german}")
        
        # Check for English text
        english_indicators = ["2025", "OK"]
        found_english = []
        for word in english_indicators:
            if word in markdown:
                found_english.append(word)
        
        print(f"English words found: {found_english}")
        
        # Assertions
        assert len(markdown) > 0, "Markdown is empty"
        assert len(found_german) > 0, f"No German text with umlauts found in markdown. Expected words: {german_words}"
        assert len(found_english) > 0, f"No English text found in markdown. Expected words: {english_indicators}"
        
        print("\n✅ CRITICAL PDF OCR test PASSED - PaddleOCR correctly extracted German+English text")
        return True
        
    except Exception as e:
        print(f"❌ CRITICAL PDF OCR test FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_plain_text_csv(token):
    """Test extraction with plain text CSV (regression test)."""
    print("\n=== TEST 9: Plain Text/CSV Regression ===")
    if not token:
        print("❌ Skipping - no auth token")
        return False
    
    try:
        # Create a simple CSV file for testing
        csv_content = """Product,Quantity,Price
Laptop,5,1200.00
Mouse,25,15.99
Keyboard,15,45.50"""
        
        files = {
            'file': ('inventory.csv', io.BytesIO(csv_content.encode('utf-8')), 'text/csv')
        }
        headers = {
            'Authorization': f'Bearer {token}'
        }
        
        response = requests.post(
            f"{BASE_URL}/api/v1/extract/sync",
            files=files,
            headers=headers,
            timeout=60
        )
        
        print(f"Status Code: {response.status_code}")
        
        if response.status_code != 200:
            print(f"Response: {response.text}")
            assert False, f"Expected 200, got {response.status_code}"
        
        data = response.json()
        result = data["result"]
        
        # Check that OCR was NOT used
        print(f"\n--- Document Processing Verification ---")
        print(f"ocrUsed: {result.get('ocrUsed')}")
        
        assert result.get("ocrUsed") == False, f"Expected ocrUsed=false for CSV, got {result.get('ocrUsed')}"
        
        # Check for required fields
        assert "markdown" in result, "No 'markdown' field in result"
        assert "structured" in result, "No 'structured' field in result"
        
        markdown = result.get("markdown", "")
        print(f"Markdown length: {len(markdown)} chars")
        
        # Verify the pipeline still works
        assert len(markdown) > 0, "Markdown is empty"
        print("✅ CSV regression test PASSED")
        return True
        
    except Exception as e:
        print(f"❌ CSV regression test FAILED: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Run all tests."""
    print("=" * 80)
    print("OmniParse Backend Regression Test Suite")
    print("Python 3.13 Migration + Dependency Upgrade Verification")
    print("numpy 2.x | opencv 5.x | onnxruntime 1.27 | pillow 12.x")
    print("=" * 80)
    print(f"Backend URL: {BASE_URL}")
    
    results = {}
    
    # Test 1: Health
    results["health"] = test_health()
    
    # Test 2: Ready
    results["ready"] = test_ready()
    
    # Test 3: Auth
    token = mint_token()
    results["auth"] = token is not None
    
    # Test 4: Image OCR - PNG
    results["image_ocr_png"] = test_image_ocr_png(token)
    
    # Test 5: Image OCR - JPG
    results["image_ocr_jpg"] = test_image_ocr_jpg(token)
    
    # Test 6: Office - DOCX
    results["office_docx"] = test_office_docx(token)
    
    # Test 7: Office - XLSX
    results["office_xlsx"] = test_office_xlsx(token)
    
    # Test 8: CRITICAL - Scanned PDF OCR
    results["scanned_pdf_ocr"] = test_ocr_german_english_pdf(token)
    
    # Test 9: Plain text/CSV regression
    results["csv_regression"] = test_plain_text_csv(token)
    
    # Summary
    print("\n" + "=" * 80)
    print("TEST SUMMARY")
    print("=" * 80)
    for test_name, passed in results.items():
        status = "✅ PASSED" if passed else "❌ FAILED"
        print(f"{test_name.upper()}: {status}")
    
    total = len(results)
    passed = sum(results.values())
    print(f"\nTotal: {passed}/{total} tests passed")
    
    if passed == total:
        print("\n🎉 ALL TESTS PASSED!")
        print("✅ Omniparse pipeline works correctly after Python 3.13 migration")
        print("✅ All dependency upgrades (numpy 2.x, opencv 5.x, onnxruntime 1.27, pillow 12.x) verified")
        return 0
    else:
        print(f"\n⚠️  {total - passed} test(s) FAILED")
        print("❌ Regression detected after dependency upgrade")
        return 1

if __name__ == "__main__":
    sys.exit(main())
